#!/usr/bin/env python3
"""Gerador de slides TCC LogPass — saida: LogPass_TCC_Slides.docx"""

from docx import Document
from docx.shared import Pt, RGBColor, Cm, Mm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── Paleta de cores (identica ao app Flutter) ────────────────────────────────
CYAN   = RGBColor(0x44, 0xCA, 0xBD)
WHITE  = RGBColor(0xE8, 0xF8, 0xF7)
GRAY   = RGBColor(0x8D, 0xDD, 0xD7)
ORANGE = RGBColor(0xFF, 0xA7, 0x26)
GREEN  = RGBColor(0x4C, 0xAF, 0x50)
PURPLE = RGBColor(0xBB, 0x86, 0xFC)

BG_HEX      = '0A1929'
SURFACE_HEX = '0D2137'
CYAN_HX     = '44CABD'
ORANGE_HX   = 'FFA726'
PURPLE_HX   = 'BB86FC'
GREEN_HX    = '4CAF50'
CARD_CYAN   = '0D3040'
CARD_ORANGE = '3B2800'
CARD_PURPLE = '2A0040'
FONT        = 'Calibri'


# ── Helpers de XML / formatacao ───────────────────────────────────────────────

def _run(para, text, color=WHITE, size=20, bold=False, italic=False):
    r = para.add_run(text)
    r.font.name    = FONT
    r.font.size    = Pt(size)
    r.font.bold    = bold
    r.font.italic  = italic
    r.font.color.rgb = color
    return r

def _para(doc_or_cell, text='', color=WHITE, size=20, bold=False, italic=False,
          align=WD_ALIGN_PARAGRAPH.LEFT, sb=0, sa=6, indent=0):
    obj = doc_or_cell if hasattr(doc_or_cell, 'add_paragraph') else doc_or_cell
    p   = obj.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(sb)
    p.paragraph_format.space_after  = Pt(sa)
    if indent:
        p.paragraph_format.left_indent = Cm(indent)
    if text:
        _run(p, text, color=color, size=size, bold=bold, italic=italic)
    return p

def slide_title(doc, text, emoji='', size=40):
    t = f'{emoji}  {text}' if emoji else text
    _para(doc, t, color=CYAN, size=size, bold=True, sb=6, sa=4)

def slide_subtitle(doc, text, color=None, size=20):
    _para(doc, text, color=color or GRAY, size=size, italic=True, sb=0, sa=10)

def hr(doc, color_hex=CYAN_HX):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(10)
    pPr   = p._p.get_or_add_pPr()
    pBdr  = OxmlElement('w:pBdr')
    btm   = OxmlElement('w:bottom')
    btm.set(qn('w:val'),   'single')
    btm.set(qn('w:sz'),    '8')
    btm.set(qn('w:space'), '1')
    btm.set(qn('w:color'), color_hex)
    pBdr.append(btm)
    pPr.append(pBdr)

def bullet(doc, text, color=WHITE, size=20, indent=1.2, sym='▸', sym_color=None):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Cm(indent)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(5)
    r1 = p.add_run(f'{sym}  ')
    r1.font.name = FONT; r1.font.size = Pt(size); r1.font.color.rgb = sym_color or CYAN
    r2 = p.add_run(text)
    r2.font.name = FONT; r2.font.size = Pt(size); r2.font.color.rgb = color

def gap(doc, n=1):
    for _ in range(n):
        _para(doc, '', size=6, sb=0, sa=0)

def pgbreak(doc):
    p    = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(0)
    run  = p.add_run()
    br   = OxmlElement('w:br')
    br.set(qn('w:type'), 'page')
    run._r.append(br)

def set_cell_bg(cell, hex_color):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    tcPr.append(shd)

def set_cell_top_border(cell, color_hex, sz='20'):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBdr = OxmlElement('w:tcBdr')
    top   = OxmlElement('w:top')
    top.set(qn('w:val'),   'single')
    top.set(qn('w:sz'),    sz)
    top.set(qn('w:space'), '0')
    top.set(qn('w:color'), color_hex)
    tcBdr.append(top)
    for side in ('left', 'bottom', 'right', 'insideH', 'insideV'):
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'),   'none')
        el.set(qn('w:sz'),    '0')
        el.set(qn('w:space'), '0')
        el.set(qn('w:color'), 'auto')
        tcBdr.append(el)
    tcPr.append(tcBdr)

def set_table_no_borders(table):
    tbl    = table._tbl
    tblPr  = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')
    tblBdr = OxmlElement('w:tblBdr')
    for side in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'),   'none')
        el.set(qn('w:sz'),    '0')
        el.set(qn('w:space'), '0')
        el.set(qn('w:color'), 'auto')
        tblBdr.append(el)
    tblPr.append(tblBdr)

def cell_para(cell, text='', color=WHITE, size=18, bold=False, italic=False,
              align=WD_ALIGN_PARAGRAPH.CENTER, sb=4, sa=4, first=False):
    p = cell.paragraphs[0] if first else cell.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(sb)
    p.paragraph_format.space_after  = Pt(sa)
    if text:
        r = p.add_run(text)
        r.font.name    = FONT
        r.font.size    = Pt(size)
        r.font.bold    = bold
        r.font.italic  = italic
        r.font.color.rgb = color
    return p


# ── Setup do documento ────────────────────────────────────────────────────────

def setup_document():
    doc     = Document()
    section = doc.sections[0]

    # A4 paisagem
    section.orientation   = WD_ORIENT.LANDSCAPE
    section.page_width    = Mm(297)
    section.page_height   = Mm(210)
    section.top_margin    = Cm(1.2)
    section.bottom_margin = Cm(1.0)
    section.left_margin   = Cm(1.8)
    section.right_margin  = Cm(1.8)

    # Fundo escuro global via XML
    bg = OxmlElement('w:background')
    bg.set(qn('w:color'), BG_HEX)
    doc.element.insert(0, bg)

    # Habilita exibicao do fundo
    settings_el = doc.settings.element
    disp_bg     = OxmlElement('w:displayBackgroundShape')
    settings_el.append(disp_bg)

    # Estilo Normal base
    normal = doc.styles['Normal']
    normal.font.name     = FONT
    normal.font.size     = Pt(18)
    normal.font.color.rgb = WHITE

    return doc


# ── Slides ────────────────────────────────────────────────────────────────────

def slide_capa(doc):
    gap(doc, 3)
    _para(doc, '💻', color=CYAN, size=56,
          align=WD_ALIGN_PARAGRAPH.CENTER, sb=0, sa=4)
    _para(doc, 'LogPass', color=CYAN, size=54, bold=True,
          align=WD_ALIGN_PARAGRAPH.CENTER, sb=0, sa=2)
    _para(doc, 'Sistema de Gerenciamento de Ocorrências',
          color=WHITE, size=24, align=WD_ALIGN_PARAGRAPH.CENTER, sb=0, sa=16)
    hr(doc)
    gap(doc)
    _para(doc, 'Apresentação de Trabalho de Conclusão de Curso',
          color=GRAY, size=18, italic=True,
          align=WD_ALIGN_PARAGRAPH.CENTER, sb=0, sa=4)
    _para(doc, 'Caua Siqueira  •  2026',
          color=GRAY, size=16, align=WD_ALIGN_PARAGRAPH.CENTER, sb=0, sa=0)


def slide_problema(doc):
    slide_title(doc, 'O Problema', '⚠️')
    slide_subtitle(doc, 'Desafios no gerenciamento de reclamações de produtos danificados')
    hr(doc)
    bullet(doc, 'Consumidores sem canal estruturado para registrar reclamações de produtos danificados', size=19)
    bullet(doc, 'Empresas sem visibilidade centralizada das ocorrências recebidas pelos clientes',        size=19)
    bullet(doc, 'Ausência de comunicação direta e rastreável entre consumidor e empresa',                 size=19)
    bullet(doc, 'Tempo de resolução sem controle — insatisfação e perda de confiança do cliente',        size=19)
    gap(doc)
    _para(doc,
          '    "Sem uma plataforma centralizada, cada ocorrência vira um problema sem dono."',
          color=GRAY, size=17, italic=True, indent=1.5)


def slide_solucao(doc):
    slide_title(doc, 'A Solução — LogPass', '✅')
    slide_subtitle(doc, 'Plataforma centralizada com 3 perfis de acesso e comunicação em tempo real')
    hr(doc)
    bullet(doc, 'Canal único para registro e acompanhamento de reclamações com protocolo exclusivo',    size=19)
    bullet(doc, 'Três perfis de acesso com funcionalidades específicas: Consumidor / Empresa / Admin',  size=19)
    bullet(doc, 'Chat em tempo real integrado — comunicação rastreável dentro da plataforma',           size=19)
    bullet(doc, 'Sistema de severidade por tempo decorrido: Aceitável (<24h) / Ruim (<48h) / Crítico', size=19)
    bullet(doc, 'Dashboard com métricas, filtros por status e análises para cada perfil',               size=19)
    gap(doc)
    _para(doc,
          '    Protocolo único por reclamação  •  Deploy em nuvem  •  App multiplataforma (Android / iOS)',
          color=GRAY, size=16, italic=True, indent=1.5)


def slide_personas_overview(doc):
    slide_title(doc, 'Personas do Sistema', '👥')
    slide_subtitle(doc, 'Três perfis com jornadas e responsabilidades distintas')
    hr(doc)
    gap(doc)

    tbl = doc.add_table(rows=1, cols=3)
    set_table_no_borders(tbl)

    personas = [
        ('👤', 'Consumidor',
         'Usuário final que registra e acompanha suas reclamações de produtos',
         CARD_CYAN, CYAN_HX),
        ('🏢', 'Empresa',
         'Gerencia ocorrências recebidas e responde diretamente aos clientes',
         CARD_ORANGE, ORANGE_HX),
        ('⚙️', 'Admin',
         'Controle total: usuários, reclamações, métricas e configurações',
         CARD_PURPLE, PURPLE_HX),
    ]

    col_width = Cm(8.5)
    for i, (ico, name, desc, bg_h, border_h) in enumerate(personas):
        cell = tbl.cell(0, i)
        cell.width = col_width
        set_cell_bg(cell, bg_h)
        set_cell_top_border(cell, border_h, sz='24')

        cell_para(cell, ico,   color=WHITE,  size=44, align=WD_ALIGN_PARAGRAPH.CENTER, sb=12, sa=4,  first=True)
        cell_para(cell, name,  color=WHITE,  size=22, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, sb=0, sa=6)
        cell_para(cell, desc,  color=GRAY,   size=15, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, sb=0, sa=14)


def slide_consumidor(doc):
    slide_title(doc, 'Persona: Consumidor', '👤', size=36)
    slide_subtitle(doc, 'João Silva — usuário final que registra reclamação de produto danificado', size=18)
    hr(doc, color_hex=CYAN_HX)
    bullet(doc, 'Registra reclamação com dados do pedido: empresa, número, motivo detalhado',    size=19)
    bullet(doc, 'Escolhe o tipo de solução desejada: Troca ou Reembolso',                        size=19)
    bullet(doc, 'Acompanha o status em tempo real: Pendente → Em Análise → Resolvida',           size=19)
    bullet(doc, 'Avalia a satisfação após a resolução da ocorrência',                             size=19)
    bullet(doc, 'Comunica-se diretamente com a empresa via chat integrado',                       size=19)
    bullet(doc, 'Recebe protocolo único de rastreamento (ex: LP1234567890)',                      size=19)
    gap(doc)
    _para(doc,
          '    Visibilidade total da ocorrência — sem depender de liguações ou e-mails',
          color=GRAY, size=16, italic=True, indent=1.5)


def slide_empresa(doc):
    slide_title(doc, 'Persona: Empresa', '🏢', size=36)
    slide_subtitle(doc, 'TechStore — empresa que gerencia reclamações dos seus clientes',
                   color=ORANGE, size=18)
    hr(doc, color_hex=ORANGE_HX)
    bullet(doc, 'Recebe notificações de reclamações pendentes dos consumidores',    size=19, sym_color=ORANGE)
    bullet(doc, 'Atualiza status em tempo real: Pendente → Em Análise → Resolvida', size=19, sym_color=ORANGE)
    bullet(doc, 'Visualiza relatórios e análises de desempenho no dashboard',        size=19, sym_color=ORANGE)
    bullet(doc, 'Meta de desempenho: resolver ocorrências em até 48h',               size=19, sym_color=ORANGE)
    bullet(doc, 'Comunica-se diretamente com o consumidor via chat integrado',       size=19, sym_color=ORANGE)
    bullet(doc, 'Recebe alertas de severidade: Aceitável / Ruim / Crítico',         size=19, sym_color=ORANGE)
    gap(doc)
    _para(doc,
          '    Mapa completo de ocorrências com filtros por status e nível de severidade',
          color=GRAY, size=16, italic=True, indent=1.5)


def slide_admin(doc):
    slide_title(doc, 'Persona: Admin', '⚙️', size=36)
    slide_subtitle(doc, 'Administrador — controle total da plataforma e dos usuários',
                   color=PURPLE, size=18)
    hr(doc, color_hex=PURPLE_HX)
    bullet(doc, 'Visualiza TODAS as reclamações do sistema com filtros por status',     size=19, sym_color=PURPLE)
    bullet(doc, 'Cria e gerencia usuários: consumidores (CPF) e empresas (CNPJ)',       size=19, sym_color=PURPLE)
    bullet(doc, 'Configura limites de severidade do sistema (horas por nível)',          size=19, sym_color=PURPLE)
    bullet(doc, 'Acessa estatísticas globais: Total / Pendentes / Em Análise / Resolvidas', size=19, sym_color=PURPLE)
    bullet(doc, 'Intervém em qualquer reclamação via chat com visão completa',           size=19, sym_color=PURPLE)
    bullet(doc, 'Controle total da plataforma com acesso exclusivo e logout seguro',    size=19, sym_color=PURPLE)
    gap(doc)
    _para(doc,
          '    Painel unificado — visão 360° de toda a operação em tempo real',
          color=GRAY, size=16, italic=True, indent=1.5)


def slide_frontend(doc):
    slide_title(doc, 'Frontend — Flutter', '📱', size=36)
    slide_subtitle(doc, 'App mobile multiplataforma com Dark Mode e animações fluidas', size=18)
    hr(doc)
    bullet(doc, 'Framework: Flutter 3 (Dart) — Android e iOS a partir de um único código',         size=19)
    bullet(doc, 'Navegação: GoRouter — rotas tipadas com redirecionamento automático por perfil',    size=19)
    bullet(doc, 'Gerenciamento de estado: Provider (ChangeNotifier) — reativo e leve',              size=19)
    bullet(doc, 'Comunicação HTTP: pacote http — integração total com a API REST do backend',        size=19)
    bullet(doc, 'Armazenamento local: SharedPreferences — token JWT persistente entre sessões',      size=19)
    bullet(doc, 'Tema: Dark Mode customizado sobre Material Design 3 — paleta Cyan #44CABD',         size=19)
    bullet(doc, 'Animações: FadeTransition, SlideTransition, ScaleTransition + partículas de fundo', size=19)


def slide_backend(doc):
    slide_title(doc, 'Backend — Node.js', '🖥️', size=36)
    slide_subtitle(doc, 'API REST robusta com autenticação JWT e deploy automatizado no Render', size=18)
    hr(doc)
    bullet(doc, 'Runtime: Node.js + Express — organizado em Controllers / Routes / Middleware',           size=19)
    bullet(doc, 'Banco de Dados: PostgreSQL hospedado no Neon Cloud — mesma instância de produção',       size=19)
    bullet(doc, 'Autenticação: JWT (JSON Web Token) — token validado por middleware em todas as rotas',   size=19)
    bullet(doc, 'Segurança: bcryptjs (hash de senhas) + express-rate-limit + express-validator',          size=19)
    bullet(doc, 'Deploy: Render — auto-deploy a cada push na branch main do repositório GitHub',          size=19)
    bullet(doc, 'Validações: CPF (consumidor) e CNPJ (empresa) com transações atômicas no PostgreSQL',   size=19)


def slide_arquitetura(doc):
    slide_title(doc, 'Arquitetura do Sistema', '🏗️', size=36)
    slide_subtitle(doc, 'Fluxo de dados entre as camadas da aplicação', size=18)
    hr(doc)
    gap(doc)

    tbl  = doc.add_table(rows=1, cols=1)
    set_table_no_borders(tbl)
    cell = tbl.cell(0, 0)
    set_cell_bg(cell, SURFACE_HEX)
    set_cell_top_border(cell, CYAN_HX, sz='16')

    arch = (
        "   📱  Flutter App  (Android / iOS)\n"
        "          │\n"
        "          │  HTTP / REST  •  JSON\n"
        "          ▼\n"
        "   🖥️   Node.js + Express  ←── JWT Middleware\n"
        "          │\n"
        "          │  SQL  (driver: pg)\n"
        "          ▼\n"
        "   🗄️   PostgreSQL  —  Neon Cloud\n\n"
        "   ☁️   Deploy: Render (backend)  •  GitHub (repositório + CI/CD)"
    )

    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(12)
    p.paragraph_format.left_indent  = Cm(1.5)
    r = p.add_run(arch)
    r.font.name      = 'Courier New'
    r.font.size      = Pt(17)
    r.font.color.rgb = CYAN


def slide_status(doc):
    slide_title(doc, 'Estado do Projeto', '📊', size=36)
    slide_subtitle(doc, 'Funcionalidades implementadas e próximos passos', size=18)
    hr(doc)

    done = [
        'Login e Registro com 3 perfis de acesso (Consumidor / Empresa / Admin)',
        'Dashboard personalizado por perfil com métricas e filtros de status',
        'Registro e acompanhamento de reclamações com protocolo único (LP...)',
        'Chat em tempo real entre consumidor e empresa — polling a cada 5s',
        'Painel Admin completo: usuários + reclamações + estatísticas globais',
        'Deploy em produção: Backend no Render + Banco PostgreSQL no Neon',
    ]
    pending = [
        'Testes end-to-end automatizados e documentação técnica final do projeto',
    ]

    for d in done:
        bullet(doc, d, size=18, sym='✅', sym_color=GREEN, color=WHITE)
    for p in pending:
        bullet(doc, p, size=18, sym='🔄', sym_color=ORANGE, color=WHITE)


def slide_conclusao(doc):
    gap(doc, 2)
    slide_title(doc, 'Conclusão', '🎯', size=42)
    hr(doc)
    gap(doc)
    bullet(doc, 'Sistema completo e funcional com 3 perfis de acesso bem definidos',      size=20)
    bullet(doc, 'Comunicação direta e rastreável entre consumidor e empresa',               size=20)
    bullet(doc, 'Visibilidade total para o administrador com painel unificado',             size=20)
    bullet(doc, 'Stack moderna e escalável: Flutter 3 + Node.js + PostgreSQL',             size=20)
    bullet(doc, 'Deploy em produção ativo — backend no Render, banco de dados no Neon',    size=20)
    gap(doc, 2)
    hr(doc)
    _para(doc,
          '"LogPass — Seus direitos protegidos em um só lugar"',
          color=CYAN, size=22, bold=True, italic=True,
          align=WD_ALIGN_PARAGRAPH.CENTER, sb=10, sa=0)


# ── Main ──────────────────────────────────────────────────────────────────────

def main():
    doc = setup_document()

    slides = [
        slide_capa,
        slide_problema,
        slide_solucao,
        slide_personas_overview,
        slide_consumidor,
        slide_empresa,
        slide_admin,
        slide_frontend,
        slide_backend,
        slide_arquitetura,
        slide_status,
        slide_conclusao,
    ]

    for i, fn in enumerate(slides):
        fn(doc)
        if i < len(slides) - 1:
            pgbreak(doc)

    output = r'C:\Users\cauas\Documents\LogPass_TCC_Slides.docx'
    doc.save(output)
    print(f'Arquivo gerado: {output}')
    print(f'Total de slides: {len(slides)}')


if __name__ == '__main__':
    main()
